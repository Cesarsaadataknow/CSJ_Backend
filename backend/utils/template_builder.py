from docx import Document

class TemplateBuilder:
    def __init__(self, template_path: str):
        self.doc = Document(template_path)

    def _insert_after(self, keyword: str, text: str):
        for i, p in enumerate(self.doc.paragraphs):
            if keyword.upper() in p.text.upper():
                self.doc.paragraphs[i + 1].text = text
                break

    def build(self, sections: dict, output_path: str):
        self._insert_after("ANTECEDENTES", sections["ANTECEDENTES"])
        self._insert_after("ANÁLISIS", sections["ANALISIS_JURIDICO"])
        self._insert_after("DECISIÓN", sections["DECISION"])

        self.doc.save(output_path)
        return output_path
