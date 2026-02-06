from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json



class DocxTemplateBuilder:
    def __init__(self, template_path: str):
        self.template_path = template_path

    def _replace_in_paragraph(self, paragraph, mapping: dict) -> None:
        if not paragraph.runs:
            return

        base_run = paragraph.runs[0]
        base_name = base_run.font.name
        base_size = base_run.font.size
        base_bold = base_run.font.bold
        base_italic = base_run.font.italic
        base_underline = base_run.font.underline

        full_text = "".join(run.text for run in paragraph.runs)
        replaced = full_text
        for k, v in mapping.items():
            token = f"{{{{{k}}}}}"
            replaced = replaced.replace(token, v or "")

        if replaced != full_text:
            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = replaced
            paragraph.runs[0].font.name = base_name or "Arial"
            paragraph.runs[0].font.size = base_size or Pt(12)
            paragraph.runs[0].font.bold = base_bold
            paragraph.runs[0].font.italic = base_italic
            paragraph.runs[0].font.underline = base_underline
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _replace_in_cell(self, cell, mapping: dict) -> None:
        for p in cell.paragraphs:
            self._replace_in_paragraph(p, mapping)

    def _replace_in_table(self, table, mapping: dict) -> None:
        for row in table.rows:
            for cell in row.cells:
                self._replace_in_cell(cell, mapping)

    def _replace_in_document(self, doc: Document, mapping: dict) -> None:
        # Body paragraphs
        for p in doc.paragraphs:
            self._replace_in_paragraph(p, mapping)

        # Body tables
        for table in doc.tables:
            self._replace_in_table(table, mapping)

        # Headers / Footers (logo suele estar aquí y NO se toca)
        for section in doc.sections:
            header = section.header
            footer = section.footer

            for p in header.paragraphs:
                self._replace_in_paragraph(p, mapping)
            for table in header.tables:
                self._replace_in_table(table, mapping)

            for p in footer.paragraphs:
                self._replace_in_paragraph(p, mapping)
            for table in footer.tables:
                self._replace_in_table(table, mapping)

    def build(self, payload: dict) -> bytes:
        doc = Document(self.template_path)

        mapping = {
            "CIUDAD_FECHA": payload.get("ciudad_fecha", ""),
            "CONSEJERO_PONENTE": payload.get("consejero_ponente", ""),
            "NUMERO_UNICO": payload.get("numero_unico", ""),
            "REFERENCIA": payload.get("referencia", ""),
            "PARTES": payload.get("partes", ""),
            "ASUNTO": payload.get("asunto", ""),
            "INTRODUCCION": payload.get("introduccion", ""),
            "ANTECEDENTES": payload.get("antecedentes", ""),
            "ACTUACION_PROCESAL": payload.get("actuacion_procesal", ""),
            "ARGUMENTOS_PARTES": payload.get("argumentos_partes", ""),
            "CONSIDERACIONES": payload.get("consideraciones", ""),
            "RECOMENDACIONES_AGENTE": payload.get("recomendaciones_agente", ""),
            "RESUELVE": payload.get("resuelve_texto", ""),  
        }

        self._replace_in_document(doc, mapping)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()


import json
from typing import Tuple, Dict, Any, List, Optional


class DocumentGeneratorService:
    """
    Genera un DOCX (bytes) usando:
    - Retrieval (userdocs o corpus)
    - LLM -> JSON estructurado según template
    - Builder -> rellena plantilla DOCX con placeholders
    """

    def __init__(
        self,
        llm_chat,
        embedder,
        indexer_userdocs,
        indexer_corpus,
        docx_builder,
        *,
        top_k_userdocs: int = 12,
        top_k_corpus: int = 12,
    ):
        self.llm_chat = llm_chat
        self.embedder = embedder
        self.indexer_userdocs = indexer_userdocs
        self.indexer_corpus = indexer_corpus
        self.docx_builder = docx_builder
        self.top_k_userdocs = top_k_userdocs
        self.top_k_corpus = top_k_corpus

    def _build_resuelve_text(self, items: List[Dict[str, Any]]) -> str:
        """
        items: [{"ordinal": "PRIMERO", "texto": "..."}, ...]
        -> "PRIMERO. ...\nSEGUNDO. ..."
        """
        lines: List[str] = []
        for it in items or []:
            ord_ = (it.get("ordinal") or "").strip().upper()
            txt = (it.get("texto") or "").strip()
            if ord_ and txt:
                lines.append(f"{ord_}. {txt}")
        return "\n".join(lines).strip()

    def _safe_json_loads(self, raw: str) -> Dict[str, Any]:
        """
        Intenta parsear JSON aunque el modelo devuelva basura antes/después.
        """
        raw = (raw or "").strip()

        # Caso ideal
        try:
            return json.loads(raw)
        except Exception:
            pass

        # Fallback: recorta del primer "{" al último "}"
        start = raw.find("{")
        end = raw.rfind("}")
        if start != -1 and end != -1 and end > start:
            try:
                return json.loads(raw[start : end + 1])
            except Exception:
                pass

        # Último fallback: devuelve vacío
        return {}

    def _detect_source(self, user_id: str, session_id: str) -> str:
        """
        Si existen docs en sesión => userdocs, si no => corpus.
        """
        try:
            files = self.indexer_userdocs.list_session_files(user_id=user_id, session_id=session_id)
            if files:
                return "userdocs"
        except Exception:
            # si falla list_session_files, no bloqueamos generación
            pass
        return "corpus"

    def _retrieve_context(
        self,
        *,
        instrucciones: str,
        user_id: str,
        session_id: str,
        source: str,
    ) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Devuelve:
        - context_str (con citas doc|chunk)
        - hits (para trazabilidad)
        """
        qvec = self.embedder.embed(instrucciones)

        if source == "userdocs":
            hits = self.indexer_userdocs.hybrid_search(
                question=instrucciones,
                query_vector=qvec,
                user_id=user_id,
                session_id=session_id,
                top_k=self.top_k_userdocs,
            )
            parts = []
            for h in hits:
                parts.append(
                    f"[{h.get('file_name')}|chunk {h.get('chunk_id')}] {h.get('content','')}"
                )
            return "\n\n".join(parts).strip(), hits

        # corpus
        hits = self.indexer_corpus.hybrid_search(
            question=instrucciones,
            query_vector=qvec,
            top_k=self.top_k_corpus,
        )
        parts = []
        for h in hits:
            # en tu corpus el texto se llama "texto" y el id de chunk es "chunk_order"
            parts.append(f"[CORPUS|chunk {h.get('chunk_order','')}] {h.get('texto','')}")
        return "\n\n".join(parts).strip(), hits

    def _build_prompt(self, *, context: str, instrucciones: str) -> str:
        """
        Prompt JSON estricto basado en el template Consejo de Estado.
        """
        return f"""
Devuelve SOLO JSON válido (sin texto extra, sin markdown, sin comillas triples).
Usa SOLO el CONTEXTO. Si falta información (ponente, número único, etc.), deja el campo vacío.

ESQUEMA:
{{
  "ciudad_fecha": "",
  "consejero_ponente": "",
  "numero_unico": "",
  "referencia": "",
  "partes": "",
  "asunto": "",
  "introduccion": "",
  "antecedentes": "",
  "actuacion_procesal": "",
  "argumentos_partes": "",
  "consideraciones": "",
  "recomendaciones_agente": "",
  "resuelve": [
    {{"ordinal": "PRIMERO", "texto": ""}},
    {{"ordinal": "SEGUNDO", "texto": ""}}
  ],
  "fuentes": [
    {{"doc": "", "chunk": ""}}
  ]
}}

CONTEXTO:
{context}

INSTRUCCIONES:
{instrucciones}
""".strip()

    # --------------------------
    # Public API
    # --------------------------
    def generate_docx_bytes(
        self,
        *,
        instrucciones: str,
        user_id: str,
        session_id: str,
        source: Optional[str] = None,   # "userdocs" | "corpus" | None (auto)
    ) -> Tuple[bytes, Dict[str, Any]]:
        """
        Retorna:
        - docx_bytes: DOCX listo
        - payload: JSON usado para llenar el template (con resuelve_texto)
        """
        instrucciones = (instrucciones or "").strip()
        if not instrucciones:
            raise ValueError("instrucciones no puede ser vacío")

        if source is None:
            source = self._detect_source(user_id=user_id, session_id=session_id)

        # 1) Retrieval
        context, hits = self._retrieve_context(
            instrucciones=instrucciones,
            user_id=user_id,
            session_id=session_id,
            source=source,
        )

        # 2) LLM -> JSON
        prompt = self._build_prompt(context=context, instrucciones=instrucciones)
        resp = self.llm_chat.invoke(prompt)

        raw = (getattr(resp, "content", None) or str(resp)).strip()
        data = self._safe_json_loads(raw)

        # Si el modelo devolvió vacío, no te rompas: genera payload mínimo
        if not data:
            data = {
                "ciudad_fecha": "",
                "consejero_ponente": "",
                "numero_unico": "",
                "referencia": "",
                "partes": "",
                "asunto": "",
                "introduccion": "",
                "antecedentes": "",
                "actuacion_procesal": "",
                "argumentos_partes": "",
                "consideraciones": "",
                "recomendaciones_agente": "No fue posible generar el contenido por fallo de formato JSON del modelo.",
                "resuelve": [],
                "fuentes": [],
            }

        # 3) Fuentes (auto) si no vienen
        if not data.get("fuentes"):
            fuentes = []
            for h in hits[:8]:
                if source == "userdocs":
                    fuentes.append({"doc": h.get("file_name", ""), "chunk": str(h.get("chunk_id", ""))})
                else:
                    fuentes.append({"doc": "CORPUS", "chunk": str(h.get("chunk_order", ""))})
            data["fuentes"] = fuentes

        # 4) resuelve_texto para el placeholder {{RESUELVE}}
        data["resuelve_texto"] = self._build_resuelve_text(data.get("resuelve", []))

        # 5) Builder DOCX (mantiene logo/estilos)
        docx_bytes = self.docx_builder.build(data)

        return docx_bytes, data

