from docx import Document
from io import BytesIO
import re
from typing import Optional
import unicodedata
from azure.identity import DefaultAzureCredential
from azure.storage.filedatalake import DataLakeServiceClient


def _replace_placeholder(container, placeholder: str, value: str):
    for paragraph in container.paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value or "")

    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_placeholder(cell, placeholder, value)


def generate_word(template_path: str, content: dict) -> bytes:
    doc = Document(template_path)

    replacements = {
        "{{ANTECEDENTES}}": content.get("antecedentes", ""),
        "{{ACTUACION_PROCESAL}}": content.get("actuacion_procesal", ""),
        "{{ARGUMENTOS_PARTES}}": content.get("argumentos_partes", ""),
        "{{CONSIDERACIONES}}": content.get("consideraciones", ""),
        "{{DECISION}}": content.get("decision", ""),
        "{{RECOMENDACIONES_IA}}": content.get("recomendaciones_ia", ""),
    }

    for placeholder, value in replacements.items():
        _replace_placeholder(doc, placeholder, value.strip())

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def upload_to_onelake(
    workspace_name: str,
    lakehouse_name: str,
    folder: str,
    filename: str,
    content_bytes: bytes
) -> str:

    service = DataLakeServiceClient(
        account_url="https://onelake.dfs.fabric.microsoft.com",
        credential=DefaultAzureCredential()
    )

    fs = service.get_file_system_client(workspace_name)
    dir_path = f"{lakehouse_name}.Lakehouse/Files/{folder}".rstrip("/")
    dir_client = fs.get_directory_client(dir_path)

    try:
        dir_client.create_directory()
    except Exception:
        pass

    file_client = dir_client.get_file_client(filename)
    file_client.upload_data(BytesIO(content_bytes), overwrite=True)

    return f"{lakehouse_name}.Lakehouse/Files/{folder}/{filename}"


def _norm(text: str) -> str:
    t = (text or "").strip().lower()
    t = unicodedata.normalize("NFD", t)
    t = "".join(ch for ch in t if unicodedata.category(ch) != "Mn")
    t = re.sub(r"\s+", " ", t)
    return t


def detect_intent(text: str) -> str:
    t = _norm(text)
    if not t:
        return "legal"

    # NEGACIÓN / CANCELAR (prioridad máxima)
    neg_patterns = [
        r"\bno\b",
        r"\bno generar\b",
        r"\bno lo hagas\b",
        r"\bno hagas\b",
        r"\bno todavia\b",
        r"\bno aun\b",
        r"\bdespues\b",
        r"\bmas tarde\b",
        r"\bcancel(a|ar|alo)\b",
        r"\bdeten\b",
        r"\bpar(a|ar)\b",
    ]
    if any(re.search(p, t) for p in neg_patterns):
        return "deny_document"

    # PRESENTACIÓN
    if re.search(r"\b(me llamo|soy|mi nombre es)\b", t):
        return "presentation"

    # GREETING
    if re.fullmatch(r"(hola|buenas|hey|buen dia|buenas tardes|buenas noches|que mas|q mas|qm)\b[!. ]*", t):
        return "greeting"

    # CAPABILITIES
    if re.search(r"\b(que haces|que puedes|ayuda|funcionalidades|capacidades|como funcionas)\b", t):
        return "capabilities"

    # CONFIRM DOCUMENT
    confirm_words = r"(si|s|dale|hagale|haga(le)?|ok|oka|okay|listo|de acuerdo|confirmo|perfecto|bien|vale|aja|melo|mandalo|envialo)"
    doc_words = r"(documento|word|docx|providencia|acta|archivo|formato|plantilla)"

    # A) Solo confirmación corta
    if re.fullmatch(rf"{confirm_words}\b[!. ]*", t):
        return "confirm_document"

    # B) Confirmación + doc
    if re.search(confirm_words, t) and re.search(doc_words, t):
        return "confirm_document"

    # C) Verbos + doc
    if re.search(r"\b(genera|generar|crea|crear|haz|hacer|arma|armar|construye|elabora|redacta|exporta)\b", t) and re.search(doc_words, t):
        return "confirm_document"

    # D) “hazlo” / “genéralo” (cuando ya preguntaste)
    if re.fullmatch(r"(hazlo|generalo|dale pues|hagale pues|listo pues|ok pues)\b[!. ]*", t):
        return "confirm_document"

    return "legal"


# =========================
# HELPERS HISTORY
# =========================
def _extract_name_from_presentation(question: str) -> Optional[str]:
    q = (question or "").strip()
    m = re.search(r"\b(me llamo|soy|mi nombre es)\s+([A-Za-zÁÉÍÓÚÜÑáéíóúüñ]+)\b", q, flags=re.IGNORECASE)
    return m.group(2).capitalize() if m else None


def _get_user_name_from_history(history: list[dict]) -> Optional[str]:
    for msg in reversed(history or []):
        name = (msg.get("extra") or {}).get("user_name")
        if name:
            return name
    return None


def _get_pending_context_from_history(history: list[dict]) -> Optional[str]:
    for msg in reversed(history or []):
        extra = msg.get("extra") or {}
        if extra.get("status") == "awaiting_confirmation" and extra.get("full_context"):
            return extra["full_context"]
    return None





# from docx import Document
# from pathlib import Path
# from io import BytesIO
# from azure.identity import DefaultAzureCredential
# from azure.storage.filedatalake import DataLakeServiceClient

# ORDER = [
#     ("I. ANTECEDENTES", "antecedentes"),
#     ("II. CONSIDERACIONES", "consideraciones"),
#     ("III. PROBLEMA JURÍDICO", "problema"),
#     ("IV. DECISIÓN", "decision"),
# ]

# def generate_word(template_path: str, content: dict)-> bytes:
#     # Cargar plantilla (solo títulos o incluso vacía)
#     doc = Document(template_path)

#     # Borrar TODO el contenido existente
#     doc._body.clear_content()

#     for title, key in ORDER:
#         # Título
#         p_title = doc.add_paragraph()
#         run = p_title.add_run(title)
#         run.bold = True

#         # Texto generado
#         p_body = doc.add_paragraph()
#         p_body.add_run(content.get(key, "").strip())

#         # Espacio entre secciones
#         doc.add_paragraph()

#     buffer = BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#     return buffer.read()



# def upload_to_onelake(
#     workspace_name: str,
#     lakehouse_name: str,
#     folder: str,
#     filename: str,
#     content_bytes: bytes) -> str:
    
#     service = DataLakeServiceClient(
#         account_url="https://onelake.dfs.fabric.microsoft.com",
#         credential=DefaultAzureCredential()
#     )

#     fs = service.get_file_system_client(workspace_name)  # filesystem = WORKSPACE
#     dir_path = f"{lakehouse_name}.Lakehouse/Files/{folder}".rstrip("/")
#     dir_client = fs.get_directory_client(dir_path)

#     # crear carpeta si no existe
#     try:
#         dir_client.create_directory()
#     except Exception:
#         pass

#     file_client = dir_client.get_file_client(filename)
#     file_client.upload_data(BytesIO(content_bytes), overwrite=True)

#     return f"{lakehouse_name}.Lakehouse/Files/{folder}/{filename}"
